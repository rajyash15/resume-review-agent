"""Document processing: turn PDF/DOCX resume files into clean text and sections.

Phase 1 of the Resume Review Agent.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

# Patterns for common resume section headings. Matched against a whole line,
# case-insensitively, so "EXPERIENCE", "Work Experience" and "skills:" all work.
SECTION_PATTERNS = {
    "summary": re.compile(
        r"^\s*(summary|professional\s+summary|profile|objective|about\s*me)\s*:?\s*$",
        re.I,
    ),
    "experience": re.compile(
        r"^\s*(professional\s+)?(experience|employment|work\s+history|work)\s*:?\s*$",
        re.I,
    ),
    "education": re.compile(
        r"^\s*(education|academic\s+background)\s*:?\s*$",
        re.I,
    ),
    "skills": re.compile(
        r"^\s*(technical\s+)?(skills|technologies?|competencies|core\s+competencies)\s*:?\s*$",
        re.I,
    ),
    "projects": re.compile(
        r"^\s*(projects|personal\s+projects|key\s+projects)\s*:?\s*$",
        re.I,
    ),
    "certifications": re.compile(
        r"^\s*(certifications?|licenses?|licenses?\s+&\s+certifications?)\s*:?\s*$",
        re.I,
    ),
    "languages": re.compile(r"^\s*(languages|additional\s+languages)\s*:?\s*$", re.I),
}

# Content appearing before any detected heading (usually name/contact info).
DEFAULT_SECTION = "header"

# Max heading length: real section headings are short; a long sentence that
# merely starts with a keyword (e.g. a summary sentence) should not match.
MAX_HEADING_LENGTH = 50


def extract_text(path: str | Path) -> str:
    """Extract raw text from a PDF or DOCX file, detecting the format by extension."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    raise ValueError(
        f"Unsupported file type: {suffix or '(none)'}. Please upload a PDF or DOCX file."
    )


def _extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            pages.append(text)
    return "\n".join(pages)


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    # Resumes occasionally use tables; include their cell text too.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def clean_text(raw: str) -> str:
    """Tidy extracted text: normalize bullets, whitespace, line breaks and empty lines."""
    if not raw:
        return ""
    # Normalize line endings (PDFs often use \r).
    text = raw.replace("\r\n", "\n").replace("\r", "\n")
    # Normalize the common bullet characters. Some PDF extractors emit a DEL
    # control char (\x7f) instead of a bullet glyph, so map that too.
    text = re.sub(r"[\u2022\u25cf\u25aa\u25e6\u00b7\u2027\u2043\x7f]", "\u2022", text)
    # Convert lines that start with a dash or asterisk bullet into real bullets.
    text = re.sub(r"^\s*-\s+", "\u2022 ", text, flags=re.M)
    text = re.sub(r"^\s*\*\s+", "\u2022 ", text, flags=re.M)
    # Collapse runs of spaces/tabs into a single space.
    text = re.sub(r"[ \t]+", " ", text)
    # Drop empty lines and strip per-line whitespace.
    lines = [line.strip() for line in text.split("\n")]
    lines = [line for line in lines if line]
    # Rejoin lines that PDF layout wrapped mid-sentence.
    lines = _reflow_wrapped_lines(lines)
    return "\n".join(lines)


# Words that typically signal a line wrapped mid-sentence ("...Python, and" /
# "Tableau."), meaning the next line continues the same sentence.
_CONTINUATION_WORDS = {
    "a", "an", "and", "as", "at", "by", "for", "from", "in", "of", "on",
    "or", "the", "to", "using", "with",
}


def _reflow_wrapped_lines(lines: list[str]) -> list[str]:
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if i + 1 < len(lines) and _looks_wrapped(line, lines[i + 1]):
            if line.endswith("-"):
                out.append(line[:-1] + lines[i + 1].strip())
            else:
                out.append(line + " " + lines[i + 1].strip())
            i += 2
        else:
            out.append(line)
            i += 1
    return out


def _looks_wrapped(line: str, next_line: str) -> bool:
    if next_line.startswith("\u2022"):  # never merge a bullet into a heading line
        return False
    if line.endswith("-"):  # hyphenated word broken across lines ("accu-racy")
        return True
    words = line.split()
    if not words:
        return False
    tail = words[-1].rstrip(".,;:").lower()
    return tail in _CONTINUATION_WORDS


def detect_sections(text: str) -> dict[str, list[str]]:
    """Split cleaned text into sections by matching heading patterns.

    Returns a mapping of section name -> list of content lines.
    """
    sections: dict[str, list[str]] = {name: [] for name in SECTION_PATTERNS}
    sections[DEFAULT_SECTION] = []
    current = DEFAULT_SECTION

    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        heading = _match_heading(line)
        if heading is not None:
            current = heading
        else:
            sections[current].append(line)

    return {name: lines for name, lines in sections.items() if lines}


def _match_heading(line: str) -> str | None:
    if len(line) > MAX_HEADING_LENGTH:
        return None
    for name, pattern in SECTION_PATTERNS.items():
        if pattern.match(line):
            return name
    return None


def parse_resume(path: str | Path) -> tuple[str, dict[str, list[str]]]:
    """Extract, clean and section a resume file.

    Returns (full_clean_text, sections).
    """
    raw = extract_text(path)
    text = clean_text(raw)
    sections = detect_sections(text)
    return text, sections
