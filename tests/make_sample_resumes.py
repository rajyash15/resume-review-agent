"""Generate realistic sample resume files (DOCX + PDF) for testing.

Dev-only script. Uses python-docx and reportlab. reportlab is not in
requirements.txt — it exists only to create the PDF test fixture.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

SAMPLE_DIR = Path(__file__).resolve().parent.parent / "sample_resumes"

# (style, text) pairs so both generators share the same content.
RESUME_CONTENT: list[tuple[str, str]] = [
    ("title", "JANE DOE"),
    ("subtitle", "Data Analyst"),
    ("contact", "jane.doe@email.com  |  (555) 123-4567  |  linkedin.com/in/janedoe"),
    ("heading", "SUMMARY"),
    ("body", "Data analyst with 4 years of experience turning raw data into "
        "actionable insights using SQL, Python, and Tableau."),
    ("heading", "EXPERIENCE"),
    ("body", "Senior Data Analyst, Acme Corp  |  2022 - Present"),
    ("bullet", "Built automated SQL pipelines reducing report generation time by 40%"),
    ("bullet", "Led a team of 3 analysts to redesign the company KPI dashboard in Tableau"),
    ("bullet", "Analyzed customer churn data, identifying trends that improved retention by 12%"),
    ("body", "Data Analyst, Beta Inc  |  2020 - 2022"),
    ("bullet", "Developed Python scripts to clean and process 10M+ rows of sales data"),
    ("bullet", "Created weekly reporting dashboards used by 15 stakeholders"),
    ("bullet", "Improved forecast accuracy by 18% with time-series models"),
    ("heading", "EDUCATION"),
    ("body", "B.S. in Statistics, State University  |  2016 - 2020"),
    ("heading", "SKILLS"),
    ("body", "SQL, Python, Tableau, Excel, Power BI, Pandas, A/B Testing, Data Visualization"),
    ("heading", "PROJECTS"),
    ("bullet", "Churn Prediction Model: logistic regression predicting customer churn with 85% accuracy"),
    ("bullet", "Sales Dashboard: interactive Tableau dashboard tracking quarterly revenue"),
]


def make_docx(path: Path) -> None:
    doc = DocxDocument()
    for style, text in RESUME_CONTENT:
        if style in ("title", "subtitle"):
            doc.add_heading(text, level=0 if style == "title" else 1)
        elif style == "contact":
            doc.add_paragraph(text)
        elif style == "heading":
            doc.add_heading(text, level=2)
        elif style == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        else:
            doc.add_paragraph(text)
    doc.save(str(path))
    print(f"Wrote {path}")


def make_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    title = ParagraphStyle("title", parent=styles["Title"], spaceAfter=2)
    subtitle = ParagraphStyle("subtitle", parent=styles["Normal"], fontName="Helvetica-Bold",
                              alignment=1, spaceAfter=6)
    contact = ParagraphStyle("contact", parent=styles["Normal"], alignment=1, spaceAfter=10)
    heading = ParagraphStyle("heading", parent=styles["Heading2"], spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("body", parent=styles["Normal"], spaceAfter=3)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=18, bulletIndent=6)

    story = []
    for style, text in RESUME_CONTENT:
        if style == "bullet":
            story.append(Paragraph(text, bullet, bulletText="\u2022"))
        elif style == "title":
            story.append(Paragraph(text, title))
        elif style == "subtitle":
            story.append(Paragraph(text, subtitle))
        elif style == "contact":
            story.append(Paragraph(text, contact))
        elif style == "heading":
            story.append(Paragraph(text, heading))
        else:
            story.append(Paragraph(text, body))
        story.append(Spacer(1, 2))

    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            leftMargin=0.75 * inch, rightMargin=0.75 * inch,
                            topMargin=0.75 * inch, bottomMargin=0.75 * inch)
    doc.build(story)
    print(f"Wrote {path}")


def main() -> None:
    SAMPLE_DIR.mkdir(exist_ok=True)
    make_docx(SAMPLE_DIR / "jane_doe_resume.docx")
    make_pdf(SAMPLE_DIR / "jane_doe_resume.pdf")


if __name__ == "__main__":
    main()
